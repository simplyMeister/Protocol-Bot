from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

def delete_table(table):
    tbl = table._tbl
    p = tbl.getparent()
    if p is not None:
        p.remove(tbl)

def fill_duty_table(doc, duty_table: list, prepared_by: str = "", day_name: str = ""):
    """
    Fill the duty tracking table in the Word document.

    Expected table columns (by index):
      0 = Role label  |  1 = Name  |  2 = Arrival time
      3 = Present/Absent  |  4 = Reason (if absent)

    Rows are matched by keyword in the first cell.
    Handles merged cells by de-duplicating on the underlying XML element.
    """
    if not doc.tables:
        return  # Template has no table – skip silently

    day_title = day_name.title()

    # ── WED-FRI Template Table Cleanup ──────────────────────────────────────
    # Table 0 is Wednesday, Table 1 is Friday. We delete the unused one.
    if len(doc.tables) > 1:
        if day_title == "Friday":
            delete_table(doc.tables[0]) # Delete Wednesday table
            table = doc.tables[0]       # Now Friday table is at index 0
        else:
            delete_table(doc.tables[1]) # Delete Friday table
            table = doc.tables[0]
    else:
        table = doc.tables[0]

    def _norm(text: str) -> str:
        # Normalize line breaks / repeated spaces from Word template labels.
        return " ".join(str(text).lower().strip().split())

    # Build lookup: normalized role -> list of duty entries
    duty_lookup: dict = {}
    for entry in duty_table:
        role = _norm(entry.get("role", ""))
        duty_lookup.setdefault(role, []).append(entry)

    for row in table.rows:
        # De-duplicate merged cells using their underlying XML element
        seen, unique_cells = set(), []
        for cell in row.cells:
            if cell._tc not in seen:
                seen.add(cell._tc)
                unique_cells.append(cell)

        if not unique_cells:
            continue

        label = _norm(unique_cells[0].text)

        # ── Provost row ───────────────────────────────────────────────────
        if "provost" in label:
            if len(unique_cells) > 1:
                unique_cells[1].paragraphs[0].text = prepared_by
            continue

        # ── Protocol role rows ────────────────────────────────────────────
        matched_entries = []
        if "entrance" in label:
            matched_entries = duty_lookup.get("entrance allocation", [])
        elif "tag allocation/collection" in label:
            matched_entries = (
                duty_lookup.get("tag allocation", [])
                + duty_lookup.get("tag collector 1", [])
                + duty_lookup.get("tag collector 2", [])
            )
        elif "tag collector 1" in label:
            matched_entries = duty_lookup.get("tag collector 1", [])
        elif "tag collector 2" in label:
            matched_entries = duty_lookup.get("tag collector 2", [])
        elif "tag allocation" in label or ("tag" in label and "allocation" in label):
            matched_entries = duty_lookup.get("tag allocation", [])
        elif "counting during" in label or ("counting" in label and "during" in label):
            if "hospi-pray" in label or "hospi pray" in label:
                matched_entries = duty_lookup.get("counting during hospi-pray", [])
            else:
                matched_entries = duty_lookup.get("counting during pre-service", [])

        if matched_entries:
            def _set(idx, value):
                if len(unique_cells) > idx:
                    unique_cells[idx].paragraphs[0].text = str(value)
            # Some templates merge tag allocation/collection into one row.
            if len(matched_entries) == 1:
                matched = matched_entries[0]
                _set(1, matched.get("name", ""))
                _set(2, matched.get("time", ""))
                _set(3, matched.get("status", ""))
                _set(4, matched.get("reason", ""))
            else:
                names = "; ".join([m.get("name", "") for m in matched_entries if m.get("name")])
                times = "; ".join([m.get("time", "") for m in matched_entries if m.get("time")])
                statuses = "; ".join([m.get("status", "") for m in matched_entries if m.get("status")])
                reasons = "; ".join([m.get("reason", "") for m in matched_entries if m.get("reason")])
                _set(1, names)
                _set(2, times)
                _set(3, statuses)
                _set(4, reasons)



def get_template_path(day_name):
    """Get the appropriate template path based on the day of the week."""
    root_dir = os.path.dirname(os.path.dirname(__file__))
    if day_name == "Saturday":
        return os.path.join(root_dir, "SATURDAY PROVOST DAILY REPORT  (School Template).docx")
    elif day_name in ["Sunday", "Tuesday", "Thursday"]:
        return os.path.join(root_dir, "SUN-TUE-THUR PROVOST DAILY REPORT  (School Template).docx")
    elif day_name in ["Friday"]:
        return os.path.join(root_dir, "WED-FRI  PROVOST DAILY REPORT  (School Template).docx")
    else:
        # Default to Sun-Tue-Thu for other days if needed
        return os.path.join(root_dir, "SUN-TUE-THUR PROVOST DAILY REPORT  (School Template).docx")

def generate_daily_report(report_data):
    """
    Generate a daily report by filling a Word template.
    """
    now = datetime.now()
    day_name = report_data.get('day_name', now.strftime("%A"))
    template_path = get_template_path(day_name)
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")
        
    doc = Document(template_path)
    
    # 1. Date Formatting
    def get_day_suffix(day):
        if 11 <= day <= 13:
            return 'TH'
        else:
            return {1: 'ST', 2: 'ND', 3: 'RD'}.get(day % 10, 'TH')

    day_num = now.day
    suffix = get_day_suffix(day_num)
    month_name = now.strftime("%B").upper()
    year = now.year
    
    # Pretty Date format (e.g. THURSDAY 28TH MAY 2026)
    pretty_date = f"{day_name.upper()} {day_num}{suffix} {month_name} {year}"
    slash_date = now.strftime("%d/%m/%Y")
    
    prepared_by = report_data.get('prepared_by', '_______________________')
    
    # Update Header Information (Date and Prepared By)
    for paragraph in doc.paragraphs:
        # Check for Attendance Summary date heading
        if "Attendance Summary" in paragraph.text:
            paragraph.text = f"Attendance Summary ({pretty_date}) {slash_date}"
            
        # Check for Date paragraph
        elif "Date –" in paragraph.text or "Date -" in paragraph.text or ("Date" in paragraph.text and "e.g" in paragraph.text):
            paragraph.text = f"• Date – ({pretty_date}) {slash_date}"
            
        # Check for Prepared by paragraph
        elif "Prepared by" in paragraph.text:
            paragraph.text = f"  • Prepared by: {prepared_by}"

            
        # Check for Friday note and clear it
        elif "On Wednesday" in paragraph.text and "Hospi pray" in paragraph.text:
            paragraph.text = ""
            
        # Check for Friday pre-service title
        elif "Communion Pre-service/Hospi pray" in paragraph.text:
            if day_name.title() == "Friday":
                paragraph.text = "Hospi pray"
            else:
                paragraph.text = "Pre-service"

    # Keep report body exactly in template layout:
    # do not prepend synthetic "Duty Assignments/Service Flow" text.

    # 3. Male / Female / Total counts (non-Saturday)
    if day_name.title() != "Saturday":
        for paragraph in doc.paragraphs:
            if paragraph.text.strip().startswith("Male:"):
                paragraph.text = f"Male: {report_data.get('male', 0)}"
            elif paragraph.text.strip().startswith("Female:"):
                paragraph.text = f"Female: {report_data.get('female', 0)}"
            elif paragraph.text.strip().startswith("Total:"):
                paragraph.text = f"Total: {report_data.get('total', 0)}"
    else:
        # Saturday specific meeting counts
        meetings = ["General Meeting", "Chaplaincy Meeting", "CHOP", "WORD FEAST"]
        for meeting in meetings:
            found = False
            meeting_key = meeting.lower().replace(" ", "_")
            counts = report_data.get(meeting_key, {"male": 0, "female": 0, "total": 0})
            
            for i, para in enumerate(doc.paragraphs):
                if meeting in para.text:
                    found = True
                    # Update next 3 lines
                    for j in range(1, 4):
                        if i + j < len(doc.paragraphs):
                            p = doc.paragraphs[i+j]
                            if "Male:" in p.text:
                                p.text = f"Male: {counts.get('male', 0)}"
                            elif "Female:" in p.text:
                                p.text = f"Female: {counts.get('female', 0)}"
                            elif "Total:" in p.text:
                                p.text = f"Total: {counts.get('total', 0)}"
                    break

    # 4. Service Overview, Challenges, Suggestions (Robust Section Replacement)
    sections = {
        "Service Overview": report_data.get('service_overview', ''),
        "Challenges & Incidents": report_data.get('challenges', ''),
        "Workflow Suggestions": report_data.get('workflow_suggestions', '')
    }

    # List of headings to mark borders
    headings = ["Service Overview", "Challenges & Incidents", "Workflow Suggestions", "REPORT DETAILS", "Attendance Summary", "Prepared by", "Male:", "Female:", "Total:", "General Meeting", "Chaplaincy Meeting", "CHOP", "WORD FEAST"]

    def is_placeholder_text(text: str) -> bool:
        lowered = text.lower().strip()
        if not lowered:
            return False
        return (
            lowered.startswith("(briefly describe")
            or lowered.startswith("example;")
            or lowered.startswith("(list any obstacles")
            or lowered.startswith("(based on today")
        )

    def replace_section(heading_text, content_text):
        found_idx = -1
        for idx, para in enumerate(doc.paragraphs):
            if heading_text in para.text:
                found_idx = idx
                break
                
        if found_idx == -1:
            return
            
        heading_para = doc.paragraphs[found_idx]
        
        # Preserve the heading line exactly as in template. Put content in the first
        # body line under that heading.
        start_clear_idx = found_idx + 1
        wrote_content = False
        if start_clear_idx < len(doc.paragraphs):
            next_p = doc.paragraphs[start_clear_idx]
            is_another_heading = any(h in next_p.text for h in headings)
            if not is_another_heading:
                next_p.text = content_text
                wrote_content = True
                start_clear_idx += 1

        if not wrote_content:
            heading_para.text = f"{heading_text}\n{content_text}"

        # Clear leftover placeholder/example lines until next heading.
        curr_idx = start_clear_idx
        while curr_idx < len(doc.paragraphs):
            p = doc.paragraphs[curr_idx]
            if any(h in p.text and len(p.text) < 60 for h in headings):
                break
            if is_placeholder_text(p.text):
                p.text = ""
            curr_idx += 1

    for heading, content in sections.items():
        replace_section(heading, content)

    # 5. Fill the duty table (SUN-TUE-THUR and WED-FRI templates only)
    if day_name.title() != "Saturday":
        duty_table = report_data.get('duty_table', [])
        fill_duty_table(doc, duty_table, prepared_by, day_name)

    # Save
    filename = f"Daily_Report_{now.strftime('%Y%m%d_%H%M%S')}.docx"
    data_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'data')
    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
    
    path = os.path.join(data_dir, filename)
    doc.save(path)
    return path

def generate_weekly_roster_report(roster_data):
    """
    Generate a weekly roster report (for future use if needed)
    """
    doc = Document()
    
    title = doc.add_heading('WEEKLY PROTOCOL ROSTER', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Generated: {roster_data.get('generated_date', datetime.now().strftime('%Y-%m-%d %H:%M'))}")
    
    roster = roster_data.get('roster', {})
    
    for service, assignments in roster.items():
        doc.add_heading(service, level=1)
        
        for role, members in assignments.items():
            doc.add_heading(role, level=2)
            if members:
                for member in members:
                    doc.add_paragraph(member, style='List Bullet')
            else:
                doc.add_paragraph("(None assigned)")
    
    # Save
    filename = f"Weekly_Roster_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    data_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'data')
    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
    
    path = os.path.join(data_dir, filename)
    doc.save(path)
    return path
