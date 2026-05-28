from docx import Document
import os

templates = [
    "SATURDAY PROVOST DAILY REPORT  (School Template).docx",
    "SUN-TUE-THUR PROVOST DAILY REPORT  (School Template).docx",
    "WED-FRI  PROVOST DAILY REPORT  (School Template).docx"
]

for template in templates:
    if os.path.exists(template):
        print(f"\n--- Inspecting: {template} ---")
        doc = Document(template)
        print("Paragraphs:")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"{i}: {para.text}")
        
        print("\nTables:")
        for i, table in enumerate(doc.tables):
            print(f"Table {i}: {len(table.rows)} rows, {len(table.columns)} columns")
            for r, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                print(f"Row {r}: {' | '.join(row_text)}")
    else:
        print(f"Template not found: {template}")
