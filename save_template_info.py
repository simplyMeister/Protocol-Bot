from docx import Document
import os

templates = [
    "SATURDAY PROVOST DAILY REPORT  (School Template).docx",
    "SUN-TUE-THUR PROVOST DAILY REPORT  (School Template).docx",
    "WED-FRI  PROVOST DAILY REPORT  (School Template).docx"
]

with open("template_structure.txt", "w", encoding="utf-8") as f:
    for template in templates:
        if os.path.exists(template):
            f.write(f"\n--- {template} ---\n")
            doc = Document(template)
            f.write("Paragraphs:\n")
            for para in doc.paragraphs:
                if para.text.strip():
                    f.write(f"- {para.text}\n")
            
            f.write("\nTables:\n")
            for i, table in enumerate(doc.tables):
                f.write(f"Table {i}:\n")
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    f.write(f"  | {' | '.join(cells)} |\n")
        else:
            f.write(f"\n--- {template} NOT FOUND ---\n")
